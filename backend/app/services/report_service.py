"""
Report generation service — uses LLM to generate analytical reports.
Exports to PDF (WeasyPrint) and DOCX (python-docx).
"""

import io
from textwrap import dedent

import markdown
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.demographics import DemographicIndicator
from app.models.municipality import Municipality
from app.models.population import PopulationRecord
from app.models.region import Region
from app.models.report import Report
from app.schemas.report import ReportGenerateRequest


class PDFExportError(RuntimeError):
    """Raised when PDF export is unavailable or fails."""


async def _gather_context(db: AsyncSession, request: ReportGenerateRequest) -> str:
    """Gather data context for LLM prompt."""
    context_parts = []

    if request.municipality_id:
        muni = await db.get(Municipality, request.municipality_id)
        if muni:
            context_parts.append(f"Муниципалитет: {muni.name}")

            pop_q = (
                select(PopulationRecord)
                .where(PopulationRecord.municipality_id == muni.id)
                .order_by(PopulationRecord.year)
            )
            pop_res = await db.execute(pop_q)
            pops = pop_res.scalars().all()
            if pops:
                pop_data = ", ".join(f"{p.year}: {p.population}" for p in pops)
                context_parts.append(f"Население по годам: {pop_data}")

            demo_q = (
                select(DemographicIndicator)
                .where(DemographicIndicator.municipality_id == muni.id)
                .order_by(DemographicIndicator.year.desc())
            )
            demo_res = await db.execute(demo_q)
            demos = demo_res.scalars().first()
            if demos:
                context_parts.append(
                    f"Демография ({demos.year}): рождаемость={demos.birth_rate}, "
                    f"смертность={demos.death_rate}, ест.прирост={demos.natural_growth_rate}, "
                    f"миграция={demos.net_migration_rate}"
                )

    elif request.region_id:
        region = await db.get(Region, request.region_id)
        if region:
            context_parts.append(f"Регион: {region.name}")

            pop_q = (
                select(PopulationRecord.year, func.sum(PopulationRecord.population))
                .join(Municipality)
                .where(Municipality.region_id == region.id)
                .group_by(PopulationRecord.year)
                .order_by(PopulationRecord.year)
            )
            pop_res = await db.execute(pop_q)
            for year, total in pop_res.all():
                context_parts.append(f"  {year}: {total}")

    return "\n".join(context_parts)


async def _call_llm(prompt: str) -> str:
    """Call the configured LLM provider."""
    from app.services.llm import get_llm

    llm = get_llm(streaming=False, temperature=0.3)
    if not llm:
        return _generate_fallback_report(prompt)

    try:
        response = await llm.ainvoke(prompt)
        return response.content
    except Exception as e:
        return _generate_fallback_report(prompt + f"\n\n(Ошибка LLM: {e})")


def _generate_fallback_report(prompt: str) -> str:
    """Generate a structured report without LLM when API key is not configured."""
    return (
        "# Аналитическая справка\n\n"
        "## Резюме\n\n"
        "Для генерации полноценной аналитической справки необходимо настроить LLM провайдер "
        "в файле `.env` (параметры `LLM_API_KEY`, `LLM_MODEL`).\n\n"
        "## Данные для анализа\n\n"
        f"```\n{prompt}\n```\n\n"
        "## Рекомендации\n\n"
        "Настройте LLM провайдер для получения AI-генерированных рекомендаций.\n"
    )


async def generate_report(db: AsyncSession, request: ReportGenerateRequest) -> Report:
    context = await _gather_context(db, request)

    prompt = (
        "Ты — аналитик демографических данных. Составь подробную аналитическую справку "
        "на русском языке в формате Markdown.\n\n"
        "Структура отчёта:\n"
        "1. Краткое резюме динамики населения\n"
        "2. Демографические тенденции и факторы влияния\n"
        "3. Прогнозная оценка на 5-10 лет\n"
        "4. Рекомендации по социальной политике и территориальному планированию\n\n"
        f"Данные:\n{context}"
    )

    content_md = await _call_llm(prompt)
    content_html = markdown.markdown(content_md, extensions=["tables", "fenced_code"])

    scope_name = "Общий"
    if request.municipality_id:
        muni = await db.get(Municipality, request.municipality_id)
        scope_name = muni.name if muni else "Муниципалитет"
    elif request.region_id:
        region = await db.get(Region, request.region_id)
        scope_name = region.name if region else "Регион"

    report = Report(
        title=f"Аналитическая справка: {scope_name}",
        report_type=request.report_type,
        scope_region_id=request.region_id,
        scope_municipality_id=request.municipality_id,
        content_markdown=content_md,
        content_html=content_html,
        parameters={
            "year_from": request.year_from,
            "year_to": request.year_to,
        },
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)
    return report


def _build_pdf_html(html_content: str | None, markdown_content: str | None = None) -> str:
    body_html = html_content
    if not body_html and markdown_content:
        body_html = markdown.markdown(markdown_content, extensions=["tables", "fenced_code"])
    if not body_html:
        body_html = "<p>Отчёт пуст.</p>"

    return dedent(
        f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="utf-8">
            <style>
                @page {{ size: A4; margin: 20mm 16mm; }}
                body {{
                    font-family: 'DejaVu Sans', Arial, sans-serif;
                    color: #1a202c;
                    line-height: 1.55;
                    font-size: 11pt;
                }}
                h1 {{
                    color: #1a365d;
                    border-bottom: 2px solid #2b6cb0;
                    padding-bottom: 10px;
                    margin-bottom: 18px;
                }}
                h2 {{
                    color: #2b6cb0;
                    margin-top: 26px;
                    margin-bottom: 10px;
                }}
                h3 {{
                    color: #2d3748;
                    margin-top: 18px;
                    margin-bottom: 8px;
                }}
                p, ul, ol {{
                    margin: 0 0 12px 0;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 15px 0;
                    font-size: 10pt;
                }}
                th, td {{
                    border: 1px solid #e2e8f0;
                    padding: 8px 10px;
                    text-align: left;
                    vertical-align: top;
                }}
                th {{
                    background-color: #edf2f7;
                    font-weight: 600;
                }}
                code {{
                    font-family: 'DejaVu Sans Mono', 'Courier New', monospace;
                    background: #f7fafc;
                    padding: 1px 4px;
                    border-radius: 4px;
                }}
                pre {{
                    white-space: pre-wrap;
                    word-break: break-word;
                    background: #f7fafc;
                    padding: 12px;
                    border-radius: 8px;
                    border: 1px solid #e2e8f0;
                }}
            </style>
        </head>
        <body>{body_html}</body>
        </html>
        """
    ).strip()


def export_pdf(html_content: str | None, markdown_content: str | None = None) -> bytes:
    """Export HTML content to PDF using WeasyPrint."""
    try:
        from weasyprint import HTML
    except Exception as exc:
        raise PDFExportError(
            "PDF export is unavailable: WeasyPrint is not installed in the backend environment."
        ) from exc

    styled_html = _build_pdf_html(html_content, markdown_content)

    try:
        pdf_bytes = HTML(string=styled_html, base_url=".").write_pdf()
    except Exception as exc:
        raise PDFExportError(f"PDF export failed: {exc}") from exc

    if not pdf_bytes or not pdf_bytes.startswith(b"%PDF"):
        raise PDFExportError("PDF export failed: generated file is not a valid PDF.")

    return pdf_bytes


def export_docx(markdown_content: str) -> bytes:
    """Export Markdown content to DOCX."""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for line in markdown_content.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
